from docx import Document
import sys

doc = Document(sys.argv[1])
parts = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        parts.append(text)
for table in doc.tables:
    for row in table.rows:
        parts.append(" | ".join(cell.text.strip() for cell in row.cells))
print("\n".join(parts))
